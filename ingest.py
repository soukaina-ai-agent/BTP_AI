"""
BTP AI - Document Ingestion Module
Handles PDF, TXT, DOCX, and image parsing, text cleaning, and chunking.
"""

import os
import re
import io
import logging
from datetime import datetime
from typing import List, Dict, Any, Optional

from dotenv import load_dotenv

load_dotenv()

logger = logging.getLogger(__name__)

IMAGE_EXTENSIONS = {"png", "jpg", "jpeg"}
OCR_LANG = os.getenv("OCR_LANG", "fra+eng")
TESSERACT_CMD = os.getenv("TESSERACT_CMD")
WINDOWS_TESSERACT_PATHS = (
    r"C:\Program Files\Tesseract-OCR\tesseract.exe",
    r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe",
)


def _resolve_tesseract_cmd() -> str | None:
    """Return an explicit Tesseract executable path when one is configured or obvious."""
    if TESSERACT_CMD:
        return TESSERACT_CMD

    if os.name == "nt":
        for candidate in WINDOWS_TESSERACT_PATHS:
            if os.path.exists(candidate):
                return candidate

    return None


class DocumentIngestor:
    """Handles loading, cleaning, and chunking of construction documents."""

    def __init__(self, chunk_size: int = 800, chunk_overlap: int = 150):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap

    # ------------------------------------------------------------------
    # Public API
    # ------------------------------------------------------------------

    def process_file(
        self, filepath: str, filename: str, project: str = "Général",
        lot: str = "", auteur: str = "", criticite: str = "Normale",
        extra_metadata: Optional[Dict[str, Any]] = None,
    ) -> List[Dict[str, Any]]:
        """
        Main entry point: load → clean → chunk → attach metadata.
        Returns a list of chunk dicts ready for embedding.
        """
        ext = filename.rsplit(".", 1)[-1].lower()

        if ext == "pdf":
            raw_text = self._extract_pdf(filepath)
        elif ext == "txt":
            raw_text = self._extract_txt(filepath)
        elif ext == "docx":
            raw_text = self._extract_docx(filepath)
        elif ext in IMAGE_EXTENSIONS:
            raw_text = self._extract_image(filepath)
        else:
            raise ValueError(f"Unsupported file type: {ext}")

        if not raw_text.strip():
            raise ValueError(f"No text could be extracted from {filename}")

        cleaned = self._clean_text(raw_text)
        chunks = self._split_text(cleaned)

        # Attach metadata to every chunk
        metadata_base = {
            "source": filename,
            "project": project,
            "lot": lot,
            "auteur": auteur,
            "criticite": criticite,
            "file_type": ext,
            "ingested_at": datetime.utcnow().isoformat(),
            "total_chunks": len(chunks),
            **(extra_metadata or {}),
        }

        result = []
        for i, chunk_text in enumerate(chunks):
            result.append({
                "text": chunk_text,
                "metadata": {**metadata_base, "chunk_index": i},
            })

        logger.info(f"[Ingestor] {filename} → {len(result)} chunks")
        return result

    # ------------------------------------------------------------------
    # Text Extraction
    # ------------------------------------------------------------------

    def _extract_pdf(self, filepath: str) -> str:
        """Extract text from PDF, falling back to OCR for scanned pages."""
        text_parts = []

        # Try pdfplumber first (better layout handling)
        try:
            import pdfplumber
            with pdfplumber.open(filepath) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
            if text_parts:
                return "\n".join(text_parts)
        except ImportError:
            pass
        except Exception as e:
            logger.warning(f"pdfplumber failed: {e}, falling back to PyPDF2")

        # Fallback: PyPDF2
        try:
            import PyPDF2
            with open(filepath, "rb") as f:
                reader = PyPDF2.PdfReader(f)
                for page in reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
        except Exception as e:
            logger.warning(f"PyPDF2 failed: {e}, falling back to PyMuPDF")

        if text_parts:
            return "\n".join(text_parts)

        # Fallback: PyMuPDF text extraction
        try:
            import fitz
            with fitz.open(filepath) as doc:
                for page in doc:
                    page_text = page.get_text("text")
                    if page_text:
                        text_parts.append(page_text)
        except ImportError:
            pass
        except Exception as e:
            logger.warning(f"PyMuPDF text extraction failed: {e}, falling back to OCR")

        if text_parts:
            return "\n".join(text_parts)

        return self._ocr_pdf(filepath)

    def _ocr_pdf(self, filepath: str) -> str:
        """Render scanned PDF pages with PyMuPDF and OCR them with Tesseract."""
        try:
            import fitz
            from PIL import Image
        except ImportError as e:
            raise RuntimeError(
                "Scanned PDF OCR requires PyMuPDF and Pillow. Install requirements.txt."
            ) from e

        text_parts = []
        try:
            with fitz.open(filepath) as doc:
                for page in doc:
                    pix = page.get_pixmap(dpi=200)
                    image = Image.open(io.BytesIO(pix.tobytes("png")))
                    ocr_text = self._ocr_image_obj(image)
                    if ocr_text:
                        text_parts.append(ocr_text)
        except Exception as e:
            raise RuntimeError(f"Failed to OCR PDF: {e}") from e

        return "\n".join(text_parts)

    def _extract_image(self, filepath: str) -> str:
        """Extract text from screenshots/photos with Tesseract OCR."""
        try:
            from PIL import Image
        except ImportError as e:
            raise RuntimeError("Image OCR requires Pillow. Install requirements.txt.") from e

        try:
            with Image.open(filepath) as image:
                return self._ocr_image_obj(image)
        except Exception as e:
            raise RuntimeError(f"Failed to OCR image: {e}") from e

    def _ocr_image_obj(self, image) -> str:
        try:
            import pytesseract
        except ImportError as e:
            raise RuntimeError("OCR requires pytesseract. Install requirements.txt.") from e

        tesseract_cmd = _resolve_tesseract_cmd()
        if tesseract_cmd:
            pytesseract.pytesseract.tesseract_cmd = tesseract_cmd

        try:
            text = pytesseract.image_to_string(image, lang=OCR_LANG)
        except pytesseract.TesseractNotFoundError as e:
            raise RuntimeError(
                "Tesseract OCR is not installed or not in PATH. Install Tesseract, "
                "or set TESSERACT_CMD in .env to the full tesseract.exe path."
            ) from e
        return text or ""

    def _extract_txt(self, filepath: str) -> str:
        """Extract text from a plain text file."""
        for encoding in ("utf-8", "latin-1", "cp1252"):
            try:
                with open(filepath, "r", encoding=encoding) as f:
                    return f.read()
            except UnicodeDecodeError:
                continue
        raise RuntimeError(f"Unable to decode text file: {filepath}")

    def _extract_docx(self, filepath: str) -> str:
        """Extract text from a DOCX file."""
        try:
            import docx
            doc = docx.Document(filepath)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            # Also grab text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(
                        cell.text.strip() for cell in row.cells if cell.text.strip()
                    )
                    if row_text:
                        paragraphs.append(row_text)
            return "\n".join(paragraphs)
        except Exception as e:
            raise RuntimeError(f"Failed to extract DOCX text: {e}")

    # ------------------------------------------------------------------
    # Text Cleaning
    # ------------------------------------------------------------------

    def _clean_text(self, text: str) -> str:
        """Remove noise and normalize whitespace."""
        # Normalize line endings
        text = text.replace("\r\n", "\n").replace("\r", "\n")
        # Remove null bytes and other control chars (keep \n and \t)
        text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]", " ", text)
        # Collapse runs of blank lines to a single blank line
        text = re.sub(r"\n{3,}", "\n\n", text)
        # Collapse horizontal whitespace (not newlines)
        text = re.sub(r"[ \t]+", " ", text)
        return text.strip()

    # ------------------------------------------------------------------
    # Text Splitting
    # ------------------------------------------------------------------

    def _split_text(self, text: str) -> List[str]:
        """
        Split text into overlapping chunks.
        Tries to split on paragraph boundaries first to preserve context.
        """
        # Split by paragraphs
        paragraphs = [p.strip() for p in text.split("\n\n") if p.strip()]

        chunks = []
        current_chunk = []
        current_length = 0

        for para in paragraphs:
            para_len = len(para)

            # If a single paragraph exceeds chunk_size, hard-split it
            if para_len > self.chunk_size:
                if current_chunk:
                    chunks.append("\n\n".join(current_chunk))
                    current_chunk, current_length = [], 0
                # Hard split the large paragraph
                for sub in self._hard_split(para):
                    chunks.append(sub)
                continue

            if current_length + para_len + 2 > self.chunk_size and current_chunk:
                chunks.append("\n\n".join(current_chunk))
                # Overlap: keep last paragraph(s) up to chunk_overlap chars
                overlap_buf, overlap_len = [], 0
                for p in reversed(current_chunk):
                    if overlap_len + len(p) < self.chunk_overlap:
                        overlap_buf.insert(0, p)
                        overlap_len += len(p)
                    else:
                        break
                current_chunk = overlap_buf
                current_length = overlap_len

            current_chunk.append(para)
            current_length += para_len + 2  # +2 for "\n\n"

        if current_chunk:
            chunks.append("\n\n".join(current_chunk))

        return [c for c in chunks if len(c.strip()) > 30]

    def _hard_split(self, text: str) -> List[str]:
        """Split a very long string by sentence or fixed length."""
        sentences = re.split(r"(?<=[.!?])\s+", text)
        chunks, current, cur_len = [], [], 0
        for sentence in sentences:
            s_len = len(sentence)
            if cur_len + s_len > self.chunk_size and current:
                chunks.append(" ".join(current))
                current, cur_len = [], 0
            current.append(sentence)
            cur_len += s_len + 1
        if current:
            chunks.append(" ".join(current))
        return chunks
